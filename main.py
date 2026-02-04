"""
DOCX/PDF Export API for GPT Actions
A simple API that generates ATS-friendly resumes and cover letters.
"""

import base64
import io
import re
from datetime import datetime
from typing import Optional, Literal
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

app = FastAPI(
    title="Document Export API",
    description="Generate ATS-friendly DOCX and PDF documents from text/markdown content. Designed for GPT Actions.",
    version="1.0.0",
    servers=[
        {"url": "https://docx-export-api.onrender.com", "description": "Production server"}
    ]
)

# CORS for GPT Actions
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://chat.openai.com", "https://chatgpt.com", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- Request/Response Models ---

class Section(BaseModel):
    """A section of the document (e.g., Experience, Education)"""
    heading: Optional[str] = Field(None, description="Section heading (e.g., 'Experience', 'Education')")
    content: str = Field(..., description="Section content as plain text or markdown")


class ExportRequest(BaseModel):
    """Request body for document export"""
    doc_type: Literal["resume", "cover_letter"] = Field(
        "resume",
        description="Type of document: 'resume' or 'cover_letter'"
    )
    file_name: Optional[str] = Field(
        None,
        description="Output filename (without extension). Defaults to 'resume' or 'cover_letter'"
    )
    title: Optional[str] = Field(
        None,
        description="Document title / applicant name (displayed at top)"
    )
    subtitle: Optional[str] = Field(
        None,
        description="Subtitle (e.g., contact info, job title)"
    )
    sections: Optional[list[Section]] = Field(
        None,
        description="Structured sections. If provided, 'content' is ignored."
    )
    content: Optional[str] = Field(
        None,
        description="Full document content as plain text or markdown. Used if 'sections' is not provided."
    )
    return_format: Literal["base64", "binary"] = Field(
        "base64",
        description="Response format: 'base64' (JSON with base64 string) or 'binary' (raw file download)"
    )

    class Config:
        json_schema_extra = {
            "example": {
                "doc_type": "resume",
                "file_name": "john_doe_resume",
                "title": "John Doe",
                "subtitle": "john.doe@email.com | (555) 123-4567 | LinkedIn: /in/johndoe",
                "sections": [
                    {"heading": "Summary", "content": "Experienced software engineer with 5+ years..."},
                    {"heading": "Experience", "content": "**Senior Developer** at TechCorp (2020-Present)\n- Led team of 5 developers\n- Increased performance by 40%"},
                    {"heading": "Education", "content": "B.S. Computer Science, State University, 2018"}
                ],
                "return_format": "base64"
            }
        }


class ExportResponse(BaseModel):
    """Response for base64 format"""
    file_name: str = Field(..., description="The generated filename with extension")
    file_base64: str = Field(..., description="Base64-encoded file content")
    mime_type: str = Field(..., description="MIME type of the file")
    message: str = Field(..., description="Success message")


# --- Document Generation ---

def parse_markdown_line(paragraph, text: str):
    """Parse simple markdown (bold, italic) and add to paragraph"""
    # Pattern for **bold** and *italic*
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'

    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # Plain text
            paragraph.add_run(match.group(4))


def create_ats_friendly_docx(request: ExportRequest) -> io.BytesIO:
    """Generate an ATS-friendly DOCX document"""
    doc = Document()

    # Set up styles for ATS compatibility (simple, standard fonts)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title (Name)
    if request.title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(request.title)
        title_run.bold = True
        title_run.font.size = Pt(18 if request.doc_type == "resume" else 14)
        title_run.font.name = 'Calibri'

    # Subtitle (Contact info)
    if request.subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(request.subtitle)
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.name = 'Calibri'

    # Add sections or content
    if request.sections:
        for section in request.sections:
            # Section heading
            if section.heading:
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(section.heading.upper())
                heading_run.bold = True
                heading_run.font.size = Pt(12)
                heading_run.font.name = 'Calibri'
                # Add a subtle line under heading
                heading_para.paragraph_format.space_after = Pt(3)

            # Section content - process line by line
            lines = section.content.strip().split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Handle bullet points
                if line.startswith('- ') or line.startswith('• '):
                    para = doc.add_paragraph(style='List Bullet')
                    parse_markdown_line(para, line[2:])
                else:
                    para = doc.add_paragraph()
                    parse_markdown_line(para, line)

                para.paragraph_format.space_after = Pt(3)

    elif request.content:
        # Parse content as markdown-ish text
        lines = request.content.strip().split('\n')
        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()  # Empty line
                continue

            # Heading detection (## Heading or HEADING:)
            if line.startswith('## '):
                para = doc.add_paragraph()
                run = para.add_run(line[3:].upper())
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith('# '):
                para = doc.add_paragraph()
                run = para.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(14)
            elif line.startswith('- ') or line.startswith('• '):
                para = doc.add_paragraph(style='List Bullet')
                parse_markdown_line(para, line[2:])
            else:
                para = doc.add_paragraph()
                parse_markdown_line(para, line)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- API Endpoints ---

@app.get("/")
async def root():
    """Health check and API info"""
    return {
        "status": "ok",
        "message": "Document Export API is running",
        "endpoints": {
            "POST /export/docx": "Generate DOCX document",
            "GET /openapi.json": "OpenAPI schema for GPT Actions"
        }
    }


@app.get("/privacy", response_class=Response)
async def privacy_policy():
    """Privacy policy for GPT Actions"""
    html = """
    <!DOCTYPE html>
    <html>
    <head><title>Privacy Policy - Document Export API</title></head>
    <body style="font-family: sans-serif; max-width: 800px; margin: 40px auto; padding: 20px;">
        <h1>Privacy Policy</h1>
        <p><strong>Last updated:</strong> February 2025</p>

        <h2>What we collect</h2>
        <p>This API processes document content (resume/cover letter text) that you submit to generate DOCX files.
        We do not store, log, or retain any of your submitted content after the document is generated and returned.</p>

        <h2>How we use your data</h2>
        <p>Your content is used solely to generate the requested document. No data is saved to any database or file system.</p>

        <h2>Third parties</h2>
        <p>We do not share any data with third parties.</p>

        <h2>Contact</h2>
        <p>For questions, contact the API administrator.</p>
    </body>
    </html>
    """
    return Response(content=html, media_type="text/html")


@app.post("/export/docx", response_model=ExportResponse, responses={
    200: {
        "description": "Successfully generated DOCX",
        "content": {
            "application/json": {
                "example": {
                    "file_name": "resume.docx",
                    "file_base64": "UEsDBBQAAAA...",
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "message": "Document generated successfully"
                }
            }
        }
    }
})
async def export_docx(request: ExportRequest):
    """
    Generate an ATS-friendly DOCX document from provided content.

    **For GPT Actions**: Use return_format='base64' and provide the base64 string
    to the user as a downloadable file.

    **Content Input** (choose one):
    - `sections`: Structured list of sections with headings and content (recommended)
    - `content`: Raw text/markdown content

    **Markdown Support**:
    - `**bold**` for bold text
    - `*italic*` for italic text
    - `- item` for bullet points
    - `## Heading` for section headings (when using content field)
    """

    if not request.sections and not request.content:
        raise HTTPException(
            status_code=400,
            detail="Either 'sections' or 'content' must be provided"
        )

    try:
        # Generate document
        docx_buffer = create_ats_friendly_docx(request)

        # Determine filename
        file_name = request.file_name or request.doc_type
        file_name = re.sub(r'[^\w\-_]', '_', file_name)  # Sanitize
        file_name = f"{file_name}.docx"

        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        if request.return_format == "binary":
            return StreamingResponse(
                docx_buffer,
                media_type=mime_type,
                headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
            )

        # Base64 response (default for GPT Actions)
        file_base64 = base64.b64encode(docx_buffer.read()).decode('utf-8')

        return ExportResponse(
            file_name=file_name,
            file_base64=file_base64,
            mime_type=mime_type,
            message="Document generated successfully. Use the base64 content to provide a download link."
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")


@app.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    """
    Generate a PDF document (converts from DOCX).

    Note: PDF generation requires additional system dependencies (LibreOffice or similar).
    For serverless deployments, consider using a third-party PDF service or
    returning DOCX and letting users convert locally.
    """
    # For v1, we'll return a helpful message about PDF limitations
    return {
        "status": "not_implemented",
        "message": "PDF export requires LibreOffice on the server. For free hosting, recommend using DOCX which users can export to PDF from Word/Google Docs.",
        "workaround": "Use /export/docx and convert to PDF locally or via Google Docs"
    }


# For local testing
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
